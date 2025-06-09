"""
Document processing utilities for extracting text from various file formats
"""

import os
import logging
import zipfile
import tempfile
import shutil
from typing import Dict, List, Optional
import PyPDF2
from docx import Document
import pandas as pd
import streamlit as st
from pathlib import Path

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Process various document formats and extract text content"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc']
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {str(e)}")
            return ""
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
            return ""
    
    def extract_text_from_doc(self, file_path: str) -> str:
        """Extract text from DOC file (legacy format)"""
        try:
            # For DOC files, we'll try to use python-docx first
            # If that fails, we'll return a message indicating manual processing needed
            try:
                return self.extract_text_from_docx(file_path)
            except:
                logger.warning(f"Could not process DOC file {file_path}. Manual processing may be required.")
                return f"[DOC file - manual processing required: {file_path}]"
        except Exception as e:
            logger.error(f"Error extracting text from DOC {file_path}: {str(e)}")
            return ""
    
    def process_single_file(self, file_path: str) -> Dict[str, str]:
        """Process a single file and return extracted text with metadata"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            return {"error": "File not found", "text": "", "filename": str(file_path)}
        
        file_extension = file_path.suffix.lower()
        filename = file_path.name
        
        if file_extension not in self.supported_formats:
            logger.warning(f"Unsupported file format: {file_extension}")
            return {"error": f"Unsupported format: {file_extension}", "text": "", "filename": filename}
        
        # Extract text based on file type
        if file_extension == '.pdf':
            text = self.extract_text_from_pdf(str(file_path))
        elif file_extension == '.docx':
            text = self.extract_text_from_docx(str(file_path))
        elif file_extension == '.doc':
            text = self.extract_text_from_doc(str(file_path))
        else:
            text = ""
        
        return {
            "filename": filename,
            "file_path": str(file_path),
            "text": text,
            "word_count": len(text.split()) if text else 0,
            "char_count": len(text) if text else 0,
            "file_size": file_path.stat().st_size,
            "extension": file_extension
        }
    
    def process_directory(self, directory_path: str) -> List[Dict[str, str]]:
        """Process all supported files in a directory"""
        directory_path = Path(directory_path)
        
        if not directory_path.exists():
            logger.error(f"Directory not found: {directory_path}")
            return []
        
        results = []
        processed_count = 0
        error_count = 0
        
        # Get all files with supported extensions
        for file_path in directory_path.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in self.supported_formats:
                logger.info(f"Processing: {file_path.name}")
                result = self.process_single_file(str(file_path))
                
                if result.get("error"):
                    error_count += 1
                    logger.error(f"Error processing {file_path.name}: {result['error']}")
                else:
                    processed_count += 1
                    logger.info(f"Successfully processed: {file_path.name} ({result['word_count']} words)")
                
                results.append(result)
        
        logger.info(f"Processing complete: {processed_count} successful, {error_count} errors")
        return results
    
    def save_extracted_text(self, results: List[Dict[str, str]], output_path: str) -> None:
        """Save extracted text results to Excel file"""
        try:
            df = pd.DataFrame(results)
            df.to_excel(output_path, index=False)
            logger.info(f"Results saved to: {output_path}")
        except Exception as e:
            logger.error(f"Error saving results: {str(e)}")
    
    def get_processing_stats(self, results: List[Dict[str, str]]) -> Dict[str, int]:
        """Get statistics about processed files"""
        stats = {
            "total_files": len(results),
            "successful": len([r for r in results if not r.get("error")]),
            "errors": len([r for r in results if r.get("error")]),
            "total_words": sum(r.get("word_count", 0) for r in results),
            "total_chars": sum(r.get("char_count", 0) for r in results),
            "pdf_files": len([r for r in results if r.get("extension") == ".pdf"]),
            "docx_files": len([r for r in results if r.get("extension") == ".docx"]),
            "doc_files": len([r for r in results if r.get("extension") == ".doc"])
        }
        return stats

    def process_uploaded_files(self, uploaded_files) -> List[Dict]:
        """Process uploaded files from Streamlit file uploader"""
        cv_data = []

        for uploaded_file in uploaded_files:
            try:
                # Handle zip files
                if uploaded_file.name.lower().endswith('.zip'):
                    cv_data.extend(self._process_zip_file(uploaded_file))
                else:
                    # Process individual file
                    result = self._process_uploaded_file(uploaded_file)
                    cv_data.append(result)
            except Exception as e:
                logger.error(f"Error processing uploaded file {uploaded_file.name}: {str(e)}")
                cv_data.append({
                    'filename': uploaded_file.name,
                    'text': '',
                    'word_count': 0,
                    'error': f"Processing error: {str(e)}"
                })

        return cv_data

    def _process_zip_file(self, zip_file) -> List[Dict]:
        """Extract and process files from a zip archive"""
        cv_data = []

        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                # Save uploaded zip file
                zip_path = os.path.join(temp_dir, zip_file.name)
                with open(zip_path, 'wb') as f:
                    f.write(zip_file.getbuffer())

                # Extract zip file
                with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                    zip_ref.extractall(temp_dir)

                # Process extracted files
                for root, dirs, files in os.walk(temp_dir):
                    for file in files:
                        if file.lower().endswith(('.pdf', '.docx', '.doc')):
                            file_path = os.path.join(root, file)
                            result = self.process_single_file(file_path)
                            cv_data.append(result)

        except Exception as e:
            logger.error(f"Error processing zip file {zip_file.name}: {str(e)}")
            cv_data.append({
                'filename': zip_file.name,
                'text': '',
                'word_count': 0,
                'error': f"Zip processing error: {str(e)}"
            })

        return cv_data

    def _process_uploaded_file(self, uploaded_file) -> Dict:
        """Process a single uploaded file"""
        try:
            file_extension = os.path.splitext(uploaded_file.name)[1].lower()

            if file_extension == '.pdf':
                text = self._extract_pdf_text_from_bytes(uploaded_file.getvalue())
            elif file_extension in ['.docx', '.doc']:
                text = self._extract_docx_text_from_bytes(uploaded_file.getvalue())
            else:
                return {
                    'filename': uploaded_file.name,
                    'text': '',
                    'word_count': 0,
                    'error': f"Unsupported file format: {file_extension}"
                }

            word_count = len(text.split()) if text else 0

            return {
                'filename': uploaded_file.name,
                'text': text,
                'word_count': word_count,
                'error': None
            }

        except Exception as e:
            logger.error(f"Error processing uploaded file {uploaded_file.name}: {str(e)}")
            return {
                'filename': uploaded_file.name,
                'text': '',
                'word_count': 0,
                'error': f"Processing error: {str(e)}"
            }

    def _extract_pdf_text_from_bytes(self, pdf_bytes) -> str:
        """Extract text from PDF bytes"""
        try:
            import io
            pdf_file = io.BytesIO(pdf_bytes)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"

            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting PDF text: {str(e)}")
            return ""

    def _extract_docx_text_from_bytes(self, docx_bytes) -> str:
        """Extract text from DOCX bytes"""
        try:
            import io
            docx_file = io.BytesIO(docx_bytes)
            doc = Document(docx_file)

            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {str(e)}")
            return ""

def main():
    """Test the document processor"""
    processor = DocumentProcessor()
    
    # Test with CV directory
    cv_directory = "CVs"
    if os.path.exists(cv_directory):
        results = processor.process_directory(cv_directory)
        stats = processor.get_processing_stats(results)
        
        print("Processing Statistics:")
        for key, value in stats.items():
            print(f"  {key}: {value}")
        
        # Save results
        processor.save_extracted_text(results, "extracted_cv_texts.xlsx")
    else:
        print(f"Directory {cv_directory} not found")

if __name__ == "__main__":
    main()
